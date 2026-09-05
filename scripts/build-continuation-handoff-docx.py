from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "pdf" / "project-continuation-handoff.docx"
PDF_OUTPUT = ROOT / "output" / "pdf" / "project-continuation-handoff.pdf"


def set_font(run, size=10, bold=False):
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.bold = bold


def paragraph(document, text, style=None):
    p = document.add_paragraph(style=style)
    run = p.add_run(text)
    set_font(run)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run)
    p.paragraph_format.space_after = Pt(3)


def heading(document, text, level=1):
    p = document.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_font(run, 15 if level == 1 else 12, True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Project Continuation Handoff")
set_font(run, 22, True)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("ai-trading-system  |  2026-09-05")
set_font(run, 10)

paragraph(doc, "Use this document to transfer the project to another ChatGPT account, Codex Cloud, Google Agent, or human contributor. The repository files named below remain authoritative when this summary conflicts with them.")

heading(doc, "Repository and Current State")
for item in [
    "GitHub: https://github.com/hamedharami-hub/ai-trading-system",
    "Default branch: main",
    "Verified handoff commit: f2628ae",
    "Immediate scope: local EURUSD / M1 / Historical Replay evidence only",
    "Phase 5 is active. Phase 6 has not started.",
]: bullet(doc, item)

heading(doc, "What Is Implemented")
for item in [
    "Deterministic Replay evidence: candles, swings, ATR-14, displacement, BOS, FVG, Order Block and Sweep/Raid facts.",
    "Immutable observation bundles, batches, JCS canonicalization and SHA-256 audit digests.",
    "Golden Dataset manifest validation, owner label cursor binding, digest verification and evidence-readiness aggregation.",
    "Offline AI foundations and a presentation-only PWA boundary.",
]: bullet(doc, item)

heading(doc, "What Is Not Implemented or Authorized")
for item in [
    "No broker, exchange or market-data network client.",
    "No Paper entry, fill, position, P and L, OMS, Demo, Testnet or Live Trading.",
    "No credentials, API keys, remote AI calls or model downloads.",
    "No claim that Golden evidence proves data quality, strategy quality or profitability.",
]: bullet(doc, item)

heading(doc, "Trading Design")
paragraph(doc, "The intended approach is deterministic market-structure and Smart Money trading with separate Scalp and Intraday engines. A future candidate requires liquidity evidence, displacement plus BOS or CHoCH, a valid OB or FVG area and invalidation, order-flow confirmation, no severe spread or volatility conflict, post-cost net R:R of at least 1.5, and a passing policy and risk boundary.")
paragraph(doc, "Definitions for OFI, CVD, exact grades, entry, stop, target and candidate scoring are not complete. Do not invent them. The locked decisions are in docs/phase-1-delegated-decisions-fa.md.")

heading(doc, "Hard Safety Boundaries")
for item in [
    "Keep LIVE_TRADING_ENABLED, BROKER_CONNECTORS_ENABLED, AI_ROUTER_ENABLED and MODEL_DOWNLOADS_ENABLED false.",
    "Unknown, stale, invalid, conflicting or timed-out state must result in no new trade.",
    "The PWA and AI cannot calculate risk, create an OrderIntent or execute a trade.",
    "Do not regard a passing build or test suite as approval for a roadmap phase or live trading.",
]: bullet(doc, item)

heading(doc, "Before Meaningful Paper Work")
for item in [
    "Use an approved local Golden Dataset with owner-labelled examples.",
    "Approve deterministic Order Flow, candidate, grade, entry, stop and target definitions.",
    "Approve evidence for cost, slippage and partial-fill assumptions.",
    "Design and approve a separate fail-closed Paper lifecycle before creating any simulated entry, fill, position or P and L.",
]: bullet(doc, item)

heading(doc, "New Agent Procedure")
for item in [
    "Read CONTINUATION-HANDOFF.md, AGENTS.md, docs/decision-register.md, docs/architecture-v2.md, docs/consolidated-roadmap-fa.md and docs/phase-1-delegated-decisions-fa.md.",
    "Use a new branch or Pull Request. Do not push directly to main without explicit owner instruction.",
    "Record a decision before material behavior changes. Add tests and preserve deterministic fail-closed behavior.",
    "Run pnpm lint, pnpm typecheck, pnpm test, pnpm build and pnpm audit --audit-level high.",
    "Report what changed, what remains prohibited, verification results and the next owner decision.",
]: bullet(doc, item)

heading(doc, "Final Review")
paragraph(doc, "After another agent opens a Pull Request, give Codex the branch name or PR URL for an architecture and safety review. Reject changes that weaken safety controls, invent trading definitions, add external access, or claim readiness without evidence.")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("ai-trading-system continuation handoff")
set_font(run, 8)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)

styles = getSampleStyleSheet()
styles.add(ParagraphStyle(name="HandoffTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=22, leading=27, spaceAfter=12, alignment=1))
styles.add(ParagraphStyle(name="HandoffHeading", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, leading=16, spaceBefore=12, spaceAfter=5))
styles.add(ParagraphStyle(name="HandoffBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.5, leading=13, spaceAfter=6))


def pdf_bullets(items):
    return ListFlowable(
        [ListItem(Paragraph(item, styles["HandoffBody"]), leftIndent=12) for item in items],
        bulletType="bullet",
        leftIndent=18,
        bulletFontName="Helvetica",
        bulletFontSize=7,
        spaceAfter=5,
    )


story = [
    Paragraph("Project Continuation Handoff", styles["HandoffTitle"]),
    Paragraph("ai-trading-system | 2026-09-05", styles["HandoffBody"]),
    Paragraph("Use this document to transfer the project to another ChatGPT account, Codex Cloud, Google Agent, or human contributor. The named repository files remain authoritative when this summary conflicts with them.", styles["HandoffBody"]),
    Paragraph("Repository and Current State", styles["HandoffHeading"]),
    pdf_bullets([
        "GitHub: https://github.com/hamedharami-hub/ai-trading-system",
        "Default branch: main. Verified handoff commit: f2628ae.",
        "Immediate scope: local EURUSD / M1 / Historical Replay evidence only.",
        "Phase 5 is active. Phase 6 has not started.",
    ]),
    Paragraph("What Is Implemented", styles["HandoffHeading"]),
    pdf_bullets([
        "Deterministic Replay evidence: candles, swings, ATR-14, displacement, BOS, FVG, Order Block and Sweep/Raid facts.",
        "Immutable observation bundles, batches, JCS canonicalization and SHA-256 audit digests.",
        "Golden Dataset manifest validation, owner label cursor binding, digest verification and evidence-readiness aggregation.",
        "Offline AI foundations and a presentation-only PWA boundary.",
    ]),
    Paragraph("What Is Not Implemented or Authorized", styles["HandoffHeading"]),
    pdf_bullets([
        "No broker, exchange or market-data network client.",
        "No Paper entry, fill, position, P and L, OMS, Demo, Testnet or Live Trading.",
        "No credentials, API keys, remote AI calls or model downloads.",
        "No claim that Golden evidence proves data quality, strategy quality or profitability.",
    ]),
    Paragraph("Trading Design", styles["HandoffHeading"]),
    Paragraph("The intended approach is deterministic market-structure and Smart Money trading with separate Scalp and Intraday engines. A future candidate requires liquidity evidence, displacement plus BOS or CHoCH, a valid OB or FVG area and invalidation, order-flow confirmation, no severe spread or volatility conflict, post-cost net R:R of at least 1.5, and a passing policy and risk boundary.", styles["HandoffBody"]),
    Paragraph("Definitions for OFI, CVD, exact grades, entry, stop, target and candidate scoring are not complete. Do not invent them. The locked decisions are in docs/phase-1-delegated-decisions-fa.md.", styles["HandoffBody"]),
    Paragraph("Hard Safety Boundaries", styles["HandoffHeading"]),
    pdf_bullets([
        "Keep LIVE_TRADING_ENABLED, BROKER_CONNECTORS_ENABLED, AI_ROUTER_ENABLED and MODEL_DOWNLOADS_ENABLED false.",
        "Unknown, stale, invalid, conflicting or timed-out state must result in no new trade.",
        "The PWA and AI cannot calculate risk, create an OrderIntent or execute a trade.",
        "Do not regard a passing build or test suite as approval for a roadmap phase or live trading.",
    ]),
    Paragraph("Before Meaningful Paper Work", styles["HandoffHeading"]),
    pdf_bullets([
        "Use an approved local Golden Dataset with owner-labelled examples.",
        "Approve deterministic Order Flow, candidate, grade, entry, stop and target definitions.",
        "Approve evidence for cost, slippage and partial-fill assumptions.",
        "Design and approve a separate fail-closed Paper lifecycle before creating any simulated entry, fill, position or P and L.",
    ]),
    Paragraph("New Agent Procedure", styles["HandoffHeading"]),
    pdf_bullets([
        "Read CONTINUATION-HANDOFF.md, AGENTS.md, docs/decision-register.md, docs/architecture-v2.md, docs/consolidated-roadmap-fa.md and docs/phase-1-delegated-decisions-fa.md.",
        "Use a new branch or Pull Request. Do not push directly to main without explicit owner instruction.",
        "Record a decision before material behavior changes. Add tests and preserve deterministic fail-closed behavior.",
        "Run pnpm lint, pnpm typecheck, pnpm test, pnpm build and pnpm audit --audit-level high.",
        "Report what changed, what remains prohibited, verification results and the next owner decision.",
    ]),
    Paragraph("Final Review", styles["HandoffHeading"]),
    Paragraph("After another agent opens a Pull Request, give Codex the branch name or PR URL for an architecture and safety review. Reject changes that weaken safety controls, invent trading definitions, add external access, or claim readiness without evidence.", styles["HandoffBody"]),
]

pdf = SimpleDocTemplate(str(PDF_OUTPUT), pagesize=A4, leftMargin=0.7 * inch, rightMargin=0.7 * inch, topMargin=0.6 * inch, bottomMargin=0.6 * inch, title="Project Continuation Handoff")
pdf.build(story)
print(OUTPUT)
print(PDF_OUTPUT)
